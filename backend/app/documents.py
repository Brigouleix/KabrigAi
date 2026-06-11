"""Création de documents PDF / Word dans Documents/Kabrig."""
import re
from pathlib import Path

OUTPUT_DIR = Path.home() / "Documents" / "Kabrig"


def _safe_name(filename: str, ext: str) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    stem = re.sub(r"[^\w\- ]", "", Path(filename).stem).strip() or "document"
    target = OUTPUT_DIR / f"{stem}{ext}"
    i = 1
    while target.exists():
        target = OUTPUT_DIR / f"{stem}-{i}{ext}"
        i += 1
    return target


def create_pdf(filename: str, title: str, content: str) -> str:
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    cell_kwargs = {"new_x": XPos.LMARGIN, "new_y": YPos.NEXT}

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    # Polices Unicode de Windows : les fonts de base de fpdf2 (latin-1 only)
    # plantent sur les tirets longs et autres caractères du LLM.
    fonts = Path("C:/Windows/Fonts")
    if (fonts / "segoeui.ttf").exists():
        pdf.add_font("ui", "", fonts / "segoeui.ttf")
        pdf.add_font("ui", "B", fonts / ("segoeuib.ttf" if (fonts / "segoeuib.ttf").exists() else "segoeui.ttf"))
        family = "ui"
    else:
        family = "helvetica"
        title = title.encode("latin-1", "replace").decode("latin-1")
        content = content.encode("latin-1", "replace").decode("latin-1")
    pdf.set_font(family, "B", 18)
    pdf.multi_cell(0, 10, title, **cell_kwargs)
    pdf.ln(4)
    pdf.set_font(family, "", 11)
    for para in content.split("\n"):
        pdf.multi_cell(0, 6, para if para.strip() else " ", **cell_kwargs)
    target = _safe_name(filename, ".pdf")
    pdf.output(str(target))
    return f"PDF créé : {target}"


def create_docx(filename: str, title: str, content: str) -> str:
    import docx

    document = docx.Document()
    document.add_heading(title, level=1)
    for para in content.split("\n"):
        document.add_paragraph(para)
    target = _safe_name(filename, ".docx")
    document.save(str(target))
    return f"Document Word créé : {target}"


def create_document(filename: str, title: str, content: str, format: str = "pdf") -> str:
    if format == "docx":
        return create_docx(filename, title, content)
    return create_pdf(filename, title, content)


DOCUMENT_TOOL_DEFINITION = {
    "type": "function",
    "function": {
        "name": "create_document",
        "description": (
            "Crée un fichier PDF ou Word (docx) dans Documents/Kabrig avec le "
            "contenu rédigé. À utiliser quand Antoine demande de produire un "
            "document, un rapport, une lettre..."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "filename": {"type": "string", "description": "Nom du fichier sans extension"},
                "title": {"type": "string", "description": "Titre affiché en haut du document"},
                "content": {"type": "string", "description": "Contenu complet du document"},
                "format": {"type": "string", "enum": ["pdf", "docx"], "description": "Défaut pdf"},
            },
            "required": ["filename", "title", "content"],
        },
    },
}
